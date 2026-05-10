#!/usr/bin/env python3
"""
APA Citation & References Consistency Checker
==============================================
Checks a thesis/paper for consistency between in-text citations
and the reference list (APA style).

Supports input: .txt, .md, .docx, .pdf
Supports output: plain text (default), .txt, .docx

Dependencies (install as needed):
    pip install python-docx   # required for .docx input AND .docx report output
    pip install pypdf          # required for .pdf input only

Usage:
    python apa_citation_checker.py thesis.docx
    python apa_citation_checker.py thesis.docx --docx report.docx
    python apa_citation_checker.py thesis.pdf  --docx report.docx
    python apa_citation_checker.py thesis.txt  --output report.txt --verbose
"""

import re
import sys
import argparse
from pathlib import Path
from dataclasses import dataclass, field
from typing import Optional


# ─────────────────────────────────────────────────────────────
#  Data structures
# ─────────────────────────────────────────────────────────────

@dataclass
class InTextCitation:
    raw: str
    authors_key: str
    year: str
    line_number: int
    context: str

@dataclass
class ReferenceEntry:
    raw: str
    authors_key: str
    year: str
    line_number: int

@dataclass
class CheckResult:
    cited_not_listed: list = field(default_factory=list)
    listed_not_cited: list = field(default_factory=list)
    matched:          list = field(default_factory=list)


# ─────────────────────────────────────────────────────────────
#  Text extraction
# ─────────────────────────────────────────────────────────────

def load_text(path: Path) -> str:
    ext = path.suffix.lower()
    if ext in (".txt", ".md"):
        return path.read_text(encoding="utf-8", errors="replace")
    elif ext == ".docx":
        try:
            from docx import Document
        except ImportError:
            sys.exit("python-docx is required for .docx files.\n"
                     "Install: pip install python-docx")
        doc = Document(str(path))
        return "\n".join(p.text for p in doc.paragraphs)
    elif ext == ".pdf":
        try:
            from pypdf import PdfReader
        except ImportError:
            try:
                from PyPDF2 import PdfReader
            except ImportError:
                sys.exit("pypdf is required for .pdf files.\n"
                         "Install: pip install pypdf")
        reader = PdfReader(str(path))
        return "\n".join(
            p.extract_text() for p in reader.pages if p.extract_text()
        )
    else:
        sys.exit(f"Unsupported file type '{ext}'. Supported: .txt .md .docx .pdf")


# ─────────────────────────────────────────────────────────────
#  Reference section detection
# ─────────────────────────────────────────────────────────────

_REF_HEADING_RE = re.compile(
    r"^\s*(?:references?|bibliography|works?\s+cited|reference\s+list)\s*$",
    re.IGNORECASE,
)
_POST_REF_RE = re.compile(
    r"^\s*(?:appendix|appendices|supplement|annex)\b",
    re.IGNORECASE,
)
_APA_ENTRY_RE = re.compile(
    r"^[A-ZÁÉÍÓÚ][a-záéíóú'\-]+.*\((?:19|20)\d{2}[a-z]?\)"
)


def split_body_and_references(text: str):
    lines = text.splitlines()
    ref_start = None

    # Strategy 1: explicit heading (take last match to skip TOC hits)
    for i, line in enumerate(lines):
        if _REF_HEADING_RE.fullmatch(line.strip()):
            ref_start = i

    # Strategy 2: heuristic — dense APA block in bottom 35 %
    if ref_start is None:
        run = 0
        first_hit = None
        for i in range(int(len(lines) * 0.65), len(lines)):
            stripped = lines[i].strip()
            if _APA_ENTRY_RE.match(stripped):
                if run == 0:
                    first_hit = i
                run += 1
                if run >= 3:
                    ref_start = first_hit
                    break
            elif stripped:
                run = 0
                first_hit = None

    if ref_start is None:
        return text, "", 0

    ref_end = len(lines)
    for i in range(ref_start + 1, len(lines)):
        if _POST_REF_RE.match(lines[i].strip()):
            ref_end = i
            break

    body = "\n".join(lines[:ref_start])
    refs  = "\n".join(lines[ref_start + 1 : ref_end])
    return body, refs, ref_start + 1   # 1-indexed


# ─────────────────────────────────────────────────────────────
#  In-text citation extraction
# ─────────────────────────────────────────────────────────────

_YEAR = r"(?:19|20)\d{2}[a-z]?|n\.d\."
_AUTH = r"[A-ZÁÉÍÓÚ][A-Za-záéíóú'\-]+"

# Parenthetical outer: (... year ...)
_PAREN_RE = re.compile(
    r"\((" + _AUTH + r"(?:\s+et\s+al\.)?[^)]*?(?:" + _YEAR + r")[^)]*?)\)",
    re.DOTALL,
)
# Single token inside parens
_TOKEN_RE = re.compile(
    r"(" + _AUTH + r"(?:\s+et\s+al\.)?)"
    r"(?:\s*[,&]\s*" + _AUTH + r"(?:\s+et\s+al\.)?)?"
    r",\s*(" + _YEAR + r")"
    r"(?:,\s*(?:pp?\.|para\.)\s*[\d\u2013\-]+)?",
)
# Narrative: Author (Year) or Author and Author (Year)
_NARR_RE = re.compile(
    r"(" + _AUTH + r"(?:\s+et\s+al\.)?)"
    r"(?:\s+and\s+" + _AUTH + r"(?:\s+et\s+al\.)?)?"
    r"\s*\((" + _YEAR + r")\)",
)


def _norm(s: str) -> str:
    return re.sub(r"\s+", " ", re.sub(r"[.,;]", "", s.lower())).strip()

def _et_al_base(k: str) -> str:
    return re.sub(r"\s+et\s+al\.?$", "", k).strip()

def _lookup_key(authors_key: str, year: str) -> str:
    return f"{_et_al_base(authors_key)}|{re.sub(r'[a-z]$', '', year)}"


def extract_citations(body: str) -> list:
    lines = body.splitlines()
    results = []
    seen: set = set()

    for lineno, line in enumerate(lines, 1):
        ctx = line.strip()[:120]

        for outer in _PAREN_RE.finditer(line):
            for seg in outer.group(1).split(";"):
                m = _TOKEN_RE.search(seg.strip())
                if m:
                    ak  = _norm(m.group(1))
                    yr  = m.group(2).strip()
                    key = _lookup_key(ak, yr)
                    if key not in seen:
                        seen.add(key)
                        results.append(InTextCitation(
                            raw=seg.strip(), authors_key=ak,
                            year=yr, line_number=lineno, context=ctx,
                        ))

        for m in _NARR_RE.finditer(line):
            ak  = _norm(m.group(1))
            yr  = m.group(2).strip()
            key = _lookup_key(ak, yr)
            if key not in seen:
                seen.add(key)
                results.append(InTextCitation(
                    raw=m.group(0), authors_key=ak,
                    year=yr, line_number=lineno, context=ctx,
                ))

    return results


# ─────────────────────────────────────────────────────────────
#  Reference list parsing
# ─────────────────────────────────────────────────────────────

_REF_ENTRY_RE = re.compile(
    r"^([A-ZÁÉÍÓÚ][A-Za-záéíóú'\-]+)[,\s].*?\((" + _YEAR + r")\)",
    re.DOTALL,
)


def parse_references(ref_text: str, start_line: int) -> list:
    lines = ref_text.splitlines()
    blocks: list = []
    cur_lines: list = []
    cur_start = 0

    for i, line in enumerate(lines):
        lineno = start_line + i
        if not line.strip():
            if cur_lines:
                blocks.append((cur_start, " ".join(cur_lines)))
                cur_lines = []
            continue
        is_new = (
            not re.match(r"^\s{4,}", line)
            and re.match(r"^[A-ZÁÉÍÓÚ\[]", line.strip())
        )
        if is_new and cur_lines:
            blocks.append((cur_start, " ".join(cur_lines)))
            cur_lines = []
        if not cur_lines:
            cur_start = lineno
        cur_lines.append(line.strip())

    if cur_lines:
        blocks.append((cur_start, " ".join(cur_lines)))

    entries = []
    for lineno, block in blocks:
        m = _REF_ENTRY_RE.match(block)
        if m:
            entries.append(ReferenceEntry(
                raw=block, authors_key=_norm(m.group(1)),
                year=m.group(2).strip(), line_number=lineno,
            ))
        else:
            ym = re.search(r"\((" + _YEAR + r")\)", block)
            fm = re.match(r"^([A-Z][A-Za-z\-]+)", block.strip())
            if ym and fm:
                entries.append(ReferenceEntry(
                    raw=block, authors_key=_norm(fm.group(1)),
                    year=ym.group(1).strip(), line_number=lineno,
                ))
    return entries


# ─────────────────────────────────────────────────────────────
#  Consistency check
# ─────────────────────────────────────────────────────────────

def check_consistency(citations: list, references: list) -> CheckResult:
    result = CheckResult()

    ref_map: dict = {}
    for r in references:
        ref_map.setdefault(_lookup_key(r.authors_key, r.year), []).append(r)

    cite_map: dict = {}
    for c in citations:
        cite_map.setdefault(_lookup_key(c.authors_key, c.year), []).append(c)

    matched_keys: set = set()

    for k, clist in cite_map.items():
        if k in ref_map:
            matched_keys.add(k)
            result.matched.append({"cite": clist[0], "ref": ref_map[k][0]})
        else:
            result.cited_not_listed.append(clist[0])

    for k, rlist in ref_map.items():
        if k not in matched_keys:
            result.listed_not_cited.append(rlist[0])

    return result


# ─────────────────────────────────────────────────────────────
#  Plain-text report
# ─────────────────────────────────────────────────────────────

SEP  = "─" * 68
SEP2 = "═" * 68

def format_text_report(result: CheckResult, citations, references,
                        ref_line: int, verbose: bool) -> str:
    out = []
    out.append(SEP2)
    out.append("  APA CITATION & REFERENCES CONSISTENCY REPORT")
    out.append(SEP2)
    out.append(f"  Unique in-text citations : {len(citations)}")
    out.append(f"  Reference entries        : {len(references)}")
    out.append(f"  Reference section line   : {ref_line}")
    out.append(f"  Matched (consistent)     : {len(result.matched)}")
    out.append(f"  Cited but not listed     : {len(result.cited_not_listed)}  <- needs fixing")
    out.append(f"  Listed but not cited     : {len(result.listed_not_cited)}  <- needs fixing")
    out.append(SEP2)

    if result.cited_not_listed:
        out.append("\n+- ISSUE 1: CITED IN TEXT - MISSING FROM REFERENCE LIST")
        out.append(SEP)
        for c in result.cited_not_listed:
            out.append(f'  Citation : "{c.raw}"')
            out.append(f"  Line     : {c.line_number}")
            out.append(f"  Context  : ...{c.context}...\n")
    else:
        out.append("\n+  All in-text citations have a reference entry.")

    if result.listed_not_cited:
        out.append("\n+- ISSUE 2: IN REFERENCE LIST - NEVER CITED IN TEXT")
        out.append(SEP)
        for r in result.listed_not_cited:
            out.append(f"  Reference : {r.raw[:110]}")
            out.append(f"  Line      : {r.line_number}\n")
    else:
        out.append("\n+  All reference entries are cited in the text.")

    if verbose and result.matched:
        out.append("\n+- MATCHED PAIRS")
        out.append(SEP)
        for m in result.matched:
            out.append(f'  + "{m["cite"].raw}"')
            out.append(f'    -> {m["ref"].raw[:90]}\n')

    out.append(SEP2)
    if not result.cited_not_listed and not result.listed_not_cited:
        out.append("  PASS - citations and references are consistent.")
    else:
        n = len(result.cited_not_listed) + len(result.listed_not_cited)
        out.append(f"  FAIL - {n} issue(s) found.")
    out.append(SEP2)
    return "\n".join(out)


# ─────────────────────────────────────────────────────────────
#  Word (.docx) report  — pure python-docx, no Node/JS
# ─────────────────────────────────────────────────────────────

def generate_docx_report(result: CheckResult, citations, references,
                          ref_line: int, filename: str, out_path: Path):
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError:
        sys.exit("python-docx is required for Word reports.\n"
                 "Install: pip install python-docx")

    # ── colour palette ──────────────────────────────────────
    C_HEAD   = RGBColor(0x1A, 0x3A, 0x5C)
    C_RED    = RGBColor(0xC0, 0x39, 0x2B)
    C_ORANGE = RGBColor(0xD3, 0x54, 0x00)
    C_GREEN  = RGBColor(0x1A, 0x7A, 0x3C)
    C_GREY   = RGBColor(0x5D, 0x6D, 0x7E)
    C_WHITE  = RGBColor(0xFF, 0xFF, 0xFF)

    BG_RED    = "FADBD8"
    BG_ORANGE = "FDEBD0"
    BG_GREEN  = "D5F5E3"
    BG_LGREY  = "F2F3F4"
    BG_HEAD   = "1A3A5C"
    BG_WHITE  = "FFFFFF"

    doc = Document()

    # ── page margins ────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1)
        section.right_margin  = Inches(1)

    # ── helper: apply both borders AND background to a cell ─
    # tcPr child order required by schema: tcW, gridSpan, vMerge,
    # tcBorders, shd, noWrap, tcMar, ...
    # We rebuild tcPr from scratch to guarantee correct ordering.
    def style_cell(cell, hex_color: str, border_color: str = "CCCCCC"):
        tc   = cell._tc
        # Remove existing tcPr and rebuild cleanly
        for existing in tc.findall(qn("w:tcPr")):
            tc.remove(existing)
        tcPr = OxmlElement("w:tcPr")
        tc.insert(0, tcPr)

        # tcBorders  (must come before shd)
        tcBorders = OxmlElement("w:tcBorders")
        for side in ("top", "left", "bottom", "right"):
            border = OxmlElement(f"w:{side}")
            border.set(qn("w:val"),   "single")
            border.set(qn("w:sz"),    "4")
            border.set(qn("w:space"), "0")
            border.set(qn("w:color"), border_color)
            tcBorders.append(border)
        tcPr.append(tcBorders)

        # shd  (must come after tcBorders)
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  hex_color)
        tcPr.append(shd)

    # Keep thin wrappers for call-site compatibility
    def set_cell_bg(cell, hex_color: str):
        style_cell(cell, hex_color)

    def set_cell_borders(cell, color="CCCCCC"):
        # Re-apply with correct ordering; bg preserved as white if not set
        tc   = cell._tc
        tcPr = tc.find(qn("w:tcPr"))
        if tcPr is None:
            style_cell(cell, "FFFFFF", color)
            return
        # Extract existing fill, then rebuild
        shd_el = tcPr.find(qn("w:shd"))
        fill   = shd_el.get(qn("w:fill"), "FFFFFF") if shd_el is not None else "FFFFFF"
        style_cell(cell, fill, color)

    # ── helper: add a run with formatting ──────────────────
    def add_run(para, text, bold=False, italic=False,
                color=None, size_pt=11):
        run = para.add_run(text)
        run.bold   = bold
        run.italic = italic
        run.font.size = Pt(size_pt)
        if color:
            run.font.color.rgb = color
        return run

    # ── helper: horizontal rule ────────────────────────────
    def add_rule(doc, color="1A3A5C"):
        p   = doc.add_paragraph()
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"),   "single")
        bottom.set(qn("w:sz"),    "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), color)
        pBdr.append(bottom)
        pPr.append(pBdr)
        p.paragraph_format.space_after = Pt(4)
        return p

    # ── helper: issue / matched table ──────────────────────
    def add_table(doc, rows_data, col_headers, col_widths_cm,
                  row_bg, text_color):
        tbl = doc.add_table(rows=1, cols=len(col_headers))
        tbl.style = "Table Grid"
        hdr = tbl.rows[0]
        for i, (hc, ht) in enumerate(zip(hdr.cells, col_headers)):
            set_cell_bg(hc, BG_HEAD)
            set_cell_borders(hc, "FFFFFF")
            hc.width = Cm(col_widths_cm[i])
            p = hc.paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            add_run(p, ht, bold=True, color=C_WHITE, size_pt=10)
        for ri, row_vals in enumerate(rows_data):
            bg  = row_bg if ri % 2 == 0 else BG_WHITE
            row = tbl.add_row()
            for ci, (cell_obj, val) in enumerate(zip(row.cells, row_vals)):
                cell_obj.width = Cm(col_widths_cm[ci])
                set_cell_bg(cell_obj, bg)
                set_cell_borders(cell_obj)
                p = cell_obj.paragraphs[0]
                p.paragraph_format.space_after = Pt(2)
                is_first = ci == 0
                add_run(p, str(val),
                        bold=is_first,
                        color=text_color if is_first else C_GREY,
                        italic=not is_first,
                        size_pt=10)
        doc.add_paragraph()

    # ════════════════════════════════════════════════════════
    #  Document content
    # ════════════════════════════════════════════════════════

    passed   = not result.cited_not_listed and not result.listed_not_cited
    n_issues = len(result.cited_not_listed) + len(result.listed_not_cited)

    # Title
    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_after = Pt(2)
    add_run(title_p, "APA Citation & References Consistency Report",
            bold=True, color=C_HEAD, size_pt=18)

    sub_p = doc.add_paragraph()
    sub_p.paragraph_format.space_after = Pt(8)
    add_run(sub_p, f"File: {filename}", color=C_GREY, size_pt=10)

    add_rule(doc)

    # Status banner
    banner_p = doc.add_paragraph()
    banner_p.paragraph_format.space_after  = Pt(12)
    banner_p.paragraph_format.space_before = Pt(6)
    if passed:
        add_run(banner_p,
                "PASS  —  Citations and references are consistent.",
                bold=True, color=C_GREEN, size_pt=12)
    else:
        add_run(banner_p,
                f"FAIL  —  {n_issues} issue(s) found. See details below.",
                bold=True, color=C_RED, size_pt=12)

    # ── Summary table ────────────────────────────────────────
    h = doc.add_heading("Summary", level=1)
    h.runs[0].font.color.rgb = C_HEAD

    tbl = doc.add_table(rows=0, cols=2)
    tbl.style = "Table Grid"

    summary_rows = [
        ("Document",             filename,
         BG_LGREY, None,    False),
        ("Reference section at", f"Line {ref_line}",
         BG_WHITE, None,    False),
        ("Unique citations",      str(len(citations)),
         BG_LGREY, None,    False),
        ("Reference entries",     str(len(references)),
         BG_WHITE, None,    False),
        ("Matched",               str(len(result.matched)),
         BG_GREEN, C_GREEN, True),
        ("Cited but not listed",
         str(len(result.cited_not_listed)),
         BG_RED    if result.cited_not_listed else BG_GREEN,
         C_RED     if result.cited_not_listed else C_GREEN, True),
        ("Listed but not cited",
         str(len(result.listed_not_cited)),
         BG_ORANGE if result.listed_not_cited else BG_GREEN,
         C_ORANGE  if result.listed_not_cited else C_GREEN, True),
        ("Overall result",
         "PASS" if passed else "FAIL",
         BG_GREEN if passed else BG_RED,
         C_GREEN  if passed else C_RED, True),
    ]
    for label, value, bg_val, val_color, bold_val in summary_rows:
        row = tbl.add_row()
        lc, vc = row.cells[0], row.cells[1]
        lc.width = Cm(7)
        vc.width = Cm(9.5)
        set_cell_bg(lc, BG_LGREY)
        set_cell_bg(vc, bg_val)
        set_cell_borders(lc)
        set_cell_borders(vc)
        lp = lc.paragraphs[0]
        vp = vc.paragraphs[0]
        lp.paragraph_format.space_after = Pt(2)
        vp.paragraph_format.space_after = Pt(2)
        add_run(lp, label, bold=True, size_pt=10.5)
        add_run(vp, value, bold=bold_val, color=val_color, size_pt=10.5)

    doc.add_paragraph()
    add_rule(doc)

    # ── Issue 1 ──────────────────────────────────────────────
    h1 = doc.add_heading(
        "Issue 1: Cited in text — missing from reference list", level=1)
    h1.runs[0].font.color.rgb = C_HEAD

    if not result.cited_not_listed:
        p = doc.add_paragraph()
        add_run(p, "None — all in-text citations have a matching reference entry.",
                bold=True, color=C_GREEN, size_pt=11)
    else:
        p = doc.add_paragraph()
        add_run(p,
                f"{len(result.cited_not_listed)} citation(s) appear in the text "
                "but have no entry in the reference list. Add the missing references.",
                color=C_RED, size_pt=11)
        doc.add_paragraph()
        add_table(
            doc,
            [(c.raw, f"Line {c.line_number}", c.context)
             for c in result.cited_not_listed],
            ["Citation (as written)", "Line", "Context snippet"],
            [5.5, 2.5, 8.5],
            BG_RED, C_RED,
        )

    add_rule(doc)

    # ── Issue 2 ──────────────────────────────────────────────
    h2 = doc.add_heading(
        "Issue 2: In reference list — never cited in text", level=1)
    h2.runs[0].font.color.rgb = C_HEAD

    if not result.listed_not_cited:
        p = doc.add_paragraph()
        add_run(p, "None — all reference entries are cited in the text.",
                bold=True, color=C_GREEN, size_pt=11)
    else:
        p = doc.add_paragraph()
        add_run(p,
                f"{len(result.listed_not_cited)} reference(s) are listed but "
                "never cited. Remove them or add missing in-text citations.",
                color=C_ORANGE, size_pt=11)
        doc.add_paragraph()
        add_table(
            doc,
            [(r.raw[:110] + ("..." if len(r.raw) > 110 else ""),
              f"Line {r.line_number}")
             for r in result.listed_not_cited],
            ["Reference entry", "Line"],
            [13.5, 3.0],
            BG_ORANGE, C_ORANGE,
        )

    add_rule(doc)

    # ── Matched pairs ────────────────────────────────────────
    h3 = doc.add_heading("Matched pairs", level=1)
    h3.runs[0].font.color.rgb = C_HEAD

    p = doc.add_paragraph()
    add_run(p,
            f"{len(result.matched)} citation(s) correctly matched "
            "to a reference entry.",
            color=C_GREEN, size_pt=11)

    if result.matched:
        doc.add_paragraph()
        add_table(
            doc,
            [(m["cite"].raw,
              m["ref"].raw[:100] + ("..." if len(m["ref"].raw) > 100 else ""))
             for m in result.matched],
            ["In-text citation", "Matched reference entry"],
            [5.5, 11.0],
            BG_GREEN, C_GREEN,
        )

    doc.save(str(out_path))

    # python-docx emits <w:zoom w:val="bestFit"/> which is invalid
    # (missing required w:percent). Patch the ZIP in-place.
    import zipfile as _zf, shutil as _sh, tempfile as _tf, os as _os
    _tmp = _tf.mkdtemp()
    try:
        with _zf.ZipFile(str(out_path), "r") as _z:
            _z.extractall(_tmp)
        _sp = _os.path.join(_tmp, "word", "settings.xml")
        if _os.path.exists(_sp):
            _xml = open(_sp, encoding="utf-8").read()
            _fixed = re.sub(
                r'<w:zoom\s+w:val="bestFit"\s*/>',
                '<w:zoom w:percent="100"/>',
                _xml,
            )
            if _fixed != _xml:
                open(_sp, "w", encoding="utf-8").write(_fixed)
        _tmp2 = str(out_path) + ".tmp"
        with _zf.ZipFile(_tmp2, "w", _zf.ZIP_DEFLATED) as _zo:
            for _r, _, _fs in _os.walk(_tmp):
                for _f in _fs:
                    _fp = _os.path.join(_r, _f)
                    _zo.write(_fp, _os.path.relpath(_fp, _tmp))
        _sh.move(_tmp2, str(out_path))
    finally:
        _sh.rmtree(_tmp, ignore_errors=True)


# ─────────────────────────────────────────────────────────────
#  Entry point
# ─────────────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(
        description="APA citation & reference list consistency checker",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog=__doc__,
    )
    parser.add_argument("file",
                        help="Thesis file (.txt, .md, .docx, .pdf)")
    parser.add_argument("-o", "--output", metavar="REPORT.txt",
                        help="Save plain-text report to this file")
    parser.add_argument("--docx", metavar="REPORT.docx",
                        help="Generate a formatted Word (.docx) report")
    parser.add_argument("-v", "--verbose", action="store_true",
                        help="Include matched pairs in plain-text output")
    args = parser.parse_args()

    path = Path(args.file)
    if not path.exists():
        sys.exit(f"File not found: {path}")

    print(f"Loading '{path.name}' ...",           file=sys.stderr)
    text = load_text(path)

    print("Locating reference section ...",        file=sys.stderr)
    body, ref_text, ref_line = split_body_and_references(text)

    if not ref_text.strip():
        print("WARNING: Could not locate a reference section.\n"
              "  Make sure your document has a heading like 'References'.",
              file=sys.stderr)

    print("Extracting in-text citations ...",      file=sys.stderr)
    citations = extract_citations(body)

    print("Parsing reference entries ...",         file=sys.stderr)
    references = parse_references(ref_text, ref_line)

    print(f"Found {len(citations)} unique citation(s), "
          f"{len(references)} reference entry/ies.",  file=sys.stderr)

    result = check_consistency(citations, references)

    # plain-text
    report = format_text_report(result, citations, references,
                                 ref_line, args.verbose)
    if args.output:
        Path(args.output).write_text(report, encoding="utf-8")
        print(f"Plain-text report -> {args.output}", file=sys.stderr)
    elif not args.docx:
        print(report)

    # Word report
    if args.docx:
        docx_path = Path(args.docx)
        print("Building Word report ...",           file=sys.stderr)
        generate_docx_report(result, citations, references,
                              ref_line, path.name, docx_path)
        print(f"Word report -> {docx_path}",        file=sys.stderr)


if __name__ == "__main__":
    main()

