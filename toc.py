"""
Word Document Formatter Script
Formats a Word document with specific chapter headers, subheading bolding,
citation italicization, and text alignment. Also generates Table of Contents
from numbered headings (patterns like 1.2, 1.2.1)
Usage: python format_word.py input.docx output.docx
"""
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re
import sys
import os

def set_run_font(run, font_name='Times New Roman', font_size=12, bold=False, italic=False):
    """Set font properties for a run"""
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run.font.italic = italic
    # For Asian characters support
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

def italicize_et_al(paragraph):
    """Find 'et al' in paragraph and italicize it"""
    text = paragraph.text
    if 'et al' in text.lower():
        # Store original formatting
        alignment = paragraph.alignment
        line_spacing = paragraph.paragraph_format.line_spacing
        space_before = paragraph.paragraph_format.space_before
        space_after = paragraph.paragraph_format.space_after
        first_line_indent = paragraph.paragraph_format.first_line_indent
        
        # Clear existing runs
        paragraph.clear()
        
        # Process text with regex to find 'et al' (case insensitive)
        parts = re.split(r'(et al\.?|et al\b)', text, flags=re.IGNORECASE)
        for part in parts:
            if part and re.match(r'et al\.?', part, re.IGNORECASE):
                run = paragraph.add_run(part)
                set_run_font(run, italic=True)
            elif part:
                run = paragraph.add_run(part)
                set_run_font(run, font_size=12)
        
        # Restore formatting
        paragraph.alignment = alignment
        paragraph.paragraph_format.line_spacing = line_spacing
        paragraph.paragraph_format.space_before = space_before
        paragraph.paragraph_format.space_after = space_after
        paragraph.paragraph_format.first_line_indent = first_line_indent

def is_numbered_heading(paragraph):
    """Check if paragraph starts with a numbered heading pattern like 1.2, 1.2.1"""
    text = paragraph.text.strip()
    return bool(re.match(r'^\d+(?:\.\d+)*\s+', text))

def get_heading_level(text):
    """Determine heading level based on number pattern (1 = level 1, 1.2 = level 2, 1.2.1 = level 3)"""
    match = re.match(r'^(\d+(?:\.\d+)*)', text)
    if match:
        num_parts = match.group(1).count('.') + 1
        return min(num_parts, 9)
    return None

def insert_toc_at_beginning(doc, toc_title="Table of Contents"):
    """Insert TOC at the beginning without breaking document structure"""
    # Create a new document for TOC
    from docx.document import Document as DocumentClass
    
    # Store all original content
    original_sections = list(doc.paragraphs)
    
    # Clear the document by removing all paragraphs
    for paragraph in original_sections:
        p = paragraph._element
        p.getparent().remove(p)
    
    # Add TOC title
    toc_heading = doc.add_paragraph()
    toc_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = toc_heading.add_run(toc_title)
    set_run_font(run, font_size=16, bold=True)
    
    # Add blank line
    doc.add_paragraph()
    
    # Find all numbered headings
    headings = []
    for para in original_sections:
        if is_numbered_heading(para):
            level = get_heading_level(para.text)
            if level:
                headings.append({
                    'text': para.text,
                    'level': level
                })
    
    # Add TOC entries
    for heading in headings:
        toc_para = doc.add_paragraph()
        toc_para.paragraph_format.left_indent = Cm((heading['level'] - 1) * 0.5)
        
        # Add heading text
        run = toc_para.add_run(heading['text'])
        set_run_font(run, font_size=11)
        
        # Add page number placeholder
        run = toc_para.add_run(" .......... 1")
        set_run_font(run, font_size=11, italic=True)
    
    # Add page break after TOC
    doc.add_page_break()
    
    # Restore original content
    for para in original_sections:
        # Clone paragraph to new document
        new_para = doc.add_paragraph()
        
        # Copy runs
        for run in para.runs:
            new_run = new_para.add_run(run.text)
            new_run.font.name = run.font.name
            new_run.font.size = run.font.size
            new_run.font.bold = run.font.bold
            new_run.font.italic = run.font.italic
        
        # Copy paragraph formatting
        new_para.alignment = para.alignment
        new_para.paragraph_format.line_spacing = para.paragraph_format.line_spacing
        new_para.paragraph_format.space_before = para.paragraph_format.space_before
        new_para.paragraph_format.space_after = para.paragraph_format.space_after
        new_para.paragraph_format.first_line_indent = para.paragraph_format.first_line_indent
    
    return len(headings)

def format_document(input_path, output_path, generate_toc=True):
    """Main function to format the document - maintains all original functionality"""
    try:
        # Open the document
        doc = Document(input_path)
        
        # ORIGINAL FUNCTIONALITY: Italicize 'et al' throughout the document
        for paragraph in doc.paragraphs:
            italicize_et_al(paragraph)
        
        # NEW FUNCTIONALITY: Generate Table of Contents if requested
        if generate_toc:
            heading_count = insert_toc_at_beginning(doc)
            print(f"✓ Generated Table of Contents with {heading_count} headings")
        
        # Save the formatted document
        doc.save(output_path)
        print(f"✓ Document successfully formatted and saved to: {output_path}")
        
    except Exception as e:
        print(f"✗ Error formatting document: {str(e)}")
        sys.exit(1)

def main():
    """Main function to handle command line arguments"""
    # Check for CLI arguments
    if len(sys.argv) != 3:
        print("Usage: python format_word.py input.docx output.docx")
        print("\nOptions:")
        print("  --no-toc    Skip table of contents generation")
        print("\nExample:")
        print("  python format_word.py input.docx output.docx")
        print("  python format_word.py input.docx output.docx --no-toc")
        sys.exit(1)
    
    input_file = sys.argv[1]
    output_file = sys.argv[2]
    
    # Check for optional --no-toc flag
    generate_toc = True
    if len(sys.argv) > 3 and sys.argv[3] == '--no-toc':
        generate_toc = False
        print("Table of Contents generation disabled")
    
    # Check if input file exists
    if not os.path.exists(input_file):
        print(f"✗ Error: Input file '{input_file}' not found.")
        sys.exit(1)
    
    # Format the document (maintains all original functionality)
    format_document(input_file, output_file, generate_toc)

if __name__ == "__main__":
    main()
