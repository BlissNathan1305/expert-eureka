"""
Word Document Formatter Script
Formats a Word document with specific chapter headers, subheading bolding, citation italicization, and text alignment
Usage: python format_word.py input.docx output.docx
"""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
import re
import sys
import os

def set_run_font(run, font_name='Times New Roman', font_size=12, bold=False, italic=False):
    """Set font properties for a run"""
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run.font.italic = italic
    # For Asian characters support
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

def italicize_et_al(paragraph):
    """Find 'et al' in paragraph and italicize it"""
    text = paragraph.text
    if 'et al' in text.lower():
        # Store original formatting
        alignment = paragraph.alignment
        line_spacing = paragraph.paragraph_format.line_spacing
        space_before = paragraph.paragraph_format.space_before
        space_after = paragraph.paragraph_format.space_after
        first_line_indent = paragraph.paragraph_format.first_line_indent
        
        # Clear existing runs
        paragraph.clear()
        
        # Process text with regex to find 'et al' (case insensitive)
        parts = re.split(r'(et al\.?|et al\b)', text, flags=re.IGNORECASE)
        
        for part in parts:
            if part and re.match(r'et al\.?', part, re.IGNORECASE):
                run = paragraph.add_run(part)
                set_run_font(run, italic=True)
            elif part:
                run = paragraph.add_run(part)
                set_run_font(run)
        
        # Restore formatting
        paragraph.alignment = alignment
        paragraph.paragraph_format.line_spacing = line_spacing
        paragraph.paragraph_format.space_before = space_before
        paragraph.paragraph_format.space_after = space_after
        paragraph.paragraph_format.first_line_indent = first_line_indent

def is_numbered_subheading(text):
    """Check if text starts with a numbered pattern like 1.1, 2.3.2, 3.4.1.2, etc."""
    # Pattern matches numbers separated by dots at the beginning
    pattern = r'^\s*(\d+\.)+\d+\s+'
    return bool(re.match(pattern, text))

def is_chapter_header(text, chapter_headers):
    """Check if text is a chapter header"""
    return text in chapter_headers

def should_indent_paragraph(paragraph_text, chapter_headers):
    """Determine if a paragraph should have first line indentation"""
    text = paragraph_text.strip()
    
    # Don't indent if:
    # 1. Empty paragraph
    if not text:
        return False
    
    # 2. Is a chapter header
    if is_chapter_header(text, chapter_headers):
        return False
    
    # 3. Is a numbered subheading
    if is_numbered_subheading(text):
        return False
    
    # Indent all other paragraphs
    return True

def format_document(input_path, output_path):
    """Main function to format the Word document"""
    
    # Check if input file exists
    if not os.path.exists(input_path):
        raise FileNotFoundError(f"Input file '{input_path}' not found")
    
    # Load the document
    print(f"Loading document: {input_path}")
    doc = Document(input_path)
    
    # Define chapter headers (exact matches)
    chapter_headers = [
        'Chapter one',
        'Introduction',
        'Chapter two',
        'Literature Review',
        'Chapter three',
        'Materials and method',
        'Chapter four',
        'Results and discussion',
        'Chapter five',
        'Conclusion and recommendations'
    ]
    
    print("Formatting document...")
    
    # Process all paragraphs
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        
        # Skip empty paragraphs but keep them (preserve spacing)
        if not text and not paragraph.runs:
            continue
        
        # 1. Center chapter headers
        if text in chapter_headers:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                set_run_font(run, bold=False)
            # Remove indentation for chapter headers
            paragraph.paragraph_format.first_line_indent = None
        
        # 2. Bold numbered subheadings
        elif is_numbered_subheading(text):
            for run in paragraph.runs:
                set_run_font(run, bold=True)
            # Remove indentation for subheadings
            paragraph.paragraph_format.first_line_indent = None
        
        # 3. Regular paragraphs - apply first line indent
        else:
            # Apply first line indent (1.27 cm = 0.5 inches, standard indent)
            paragraph.paragraph_format.first_line_indent = Cm(1.27)
        
        # 4. Set text alignment to justified (straight at both ends)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        # 5. Set line spacing to double
        paragraph.paragraph_format.line_spacing = WD_LINE_SPACING.DOUBLE
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.space_before = Pt(0)
        
        # 6. Apply Times New Roman, 12pt to all runs
        for run in paragraph.runs:
            current_bold = run.font.bold
            current_italic = run.font.italic
            set_run_font(run, bold=current_bold, italic=current_italic)
        
        # 7. Italicize 'et al' (do this after other formatting)
        if text and 'et al' in text.lower():
            italicize_et_al(paragraph)
    
    # Process tables if present
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    text = paragraph.text.strip()
                    
                    if text in chapter_headers:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        paragraph.paragraph_format.first_line_indent = None
                        for run in paragraph.runs:
                            set_run_font(run, bold=False)
                    
                    elif is_numbered_subheading(text):
                        for run in paragraph.runs:
                            set_run_font(run, bold=True)
                        paragraph.paragraph_format.first_line_indent = None
                    
                    else:
                        # Apply first line indent for regular text in tables
                        paragraph.paragraph_format.first_line_indent = Cm(1.27)
                    
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    paragraph.paragraph_format.line_spacing = WD_LINE_SPACING.DOUBLE
                    paragraph.paragraph_format.space_after = Pt(0)
                    paragraph.paragraph_format.space_before = Pt(0)
                    
                    for run in paragraph.runs:
                        current_bold = run.font.bold
                        current_italic = run.font.italic
                        set_run_font(run, bold=current_bold, italic=current_italic)
                    
                    if text and 'et al' in text.lower():
                        italicize_et_al(paragraph)
    
    # Save the formatted document
    doc.save(output_path)
    print(f"✓ Document formatted successfully!")
    print(f"  Saved to: {output_path}")
    print(f"  Formatting applied:")
    print(f"    • Chapter headers: centered")
    print(f"    • Numbered subheadings: bold, no indent")
    print(f"    • Regular paragraphs: first line indent (1.27 cm)")
    print(f"    • 'et al' citations: italicized")
    print(f"    • Text alignment: justified")
    print(f"    • Font: Times New Roman, 12pt")
    print(f"    • Line spacing: double")

def print_usage():
    """Print usage instructions"""
    print("=" * 60)
    print("WORD DOCUMENT FORMATTER")
    print("=" * 60)
    print("Usage: python format_word.py <input_file> <output_file>")
    print("\nExamples:")
    print("  python format_word.py my_thesis.docx formatted_thesis.docx")
    print("  python format_word.py input.docx output.docx")
    print("\nArguments:")
    print("  input_file   : Path to the Word document to format")
    print("  output_file  : Path where the formatted document will be saved")
    print("\nFormatting applied:")
    print("  • Chapter headers centered")
    print("  • Numbered subheadings (e.g., 1.1, 2.3.2) made bold")
    print("  • Regular paragraphs have first line indent (1.27 cm)")
    print("  • 'et al' in citations italicized")
    print("  • Text justified (straight at both ends)")
    print("  • Times New Roman, 12pt font")
    print("  • Double line spacing")
    print("=" * 60)

if __name__ == "__main__":
    # Check command line arguments
    if len(sys.argv) != 3:
        print_usage()
        sys.exit(1)
    
    input_file = sys.argv[1]
    output_file = sys.argv[2]
    
    try:
        format_document(input_file, output_file)
    except FileNotFoundError as e:
        print(f"\n✗ Error: {e}")
        sys.exit(1)
    except Exception as e:
        print(f"\n✗ An error occurred: {str(e)}")
        print("\nTroubleshooting tips:")
        print("  • Make sure the input file is not open in another program")
        print("  • Check that you have write permissions for the output location")
        print("  • Verify the input file is a valid .docx file")
        sys.exit(1)
