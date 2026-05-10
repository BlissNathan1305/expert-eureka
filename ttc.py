from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re
import sys
import os
from docx import Document

def set_run_font(run, font_name='Times New Roman', font_size=12, bold=False, italic=False):
    """Set font properties for a run"""
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run.font.italic = italic
    # For Asian characters support
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

def italicize_et_al(paragraph):
    """Find 'et al' in paragraph and italicize it"""
    text = paragraph.text
    if 'et al' in text.lower():
        # Store original formatting
        alignment = paragraph.alignment
        line_spacing = paragraph.paragraph_format.line_spacing
        space_before = paragraph.paragraph_format.space_before
        space_after = paragraph.paragraph_format.space_after
        first_line_indent = paragraph.paragraph_format.first_line_indent
        
        # Clear existing runs
        paragraph.clear()
        
        # Process text with regex to find 'et al' (case insensitive)
        parts = re.split(r'(et al\.?|et al\b)', text, flags=re.IGNORECASE)
        for part in parts:
            if part and re.match(r'et al\.?', part, re.IGNORECASE):
                run = paragraph.add_run(part)
                set_run_font(run, italic=True)
            elif part:
                run = paragraph.add_run(part)
                set_run_font(run, font_size=12)
        
        # Restore formatting
        paragraph.alignment = alignment
        paragraph.paragraph_format.line_spacing = line_spacing
        paragraph.paragraph_format.space_before = space_before
        paragraph.paragraph_format.space_after = space_after
        paragraph.paragraph_format.first_line_indent = first_line_indent

def is_numbered_heading(paragraph):
    """Check if paragraph starts with a numbered heading pattern like 1.2, 1.2.1"""
    text = paragraph.text.strip()
    return bool(re.match(r'^\d+(?:\.\d+)*\s+', text))

def get_heading_level(text):
    """Determine heading level based on number pattern (1 = level 1, 1.2 = level 2, 1.2.1 = level 3)"""
    match = re.match(r'^(\d+(?:\.\d+)*)', text)
    if match:
        num_parts = match.group(1).count('.') + 1
        return min(num_parts, 9)
    return None

def add_page_numbers(doc):
    """Add page numbers to the bottom center of every page"""
    for section in doc.sections:
        footer = section.footer
        footer.paragraphs[0].clear()
        
        footer_para = footer.paragraphs[0]
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add "Page " text
        run = footer_para.add_run("Page ")
        set_run_font(run, font_size=10)
        
        # Add PAGE field code
        run = footer_para.add_run()
        fldChar = OxmlElement('w:fldChar')
        fldChar.set(qn('w:fldCharType'), 'begin')
        run._element.append(fldChar)
        
        instrText = OxmlElement('w:instrText')
        instrText.text = "PAGE"
        run._element.append(instrText)
        
        fldChar = OxmlElement('w:fldChar')
        fldChar.set(qn('w:fldCharType'), 'end')
        run._element.append(fldChar)
        
        set_run_font(run, font_size=10)

def generate_toc_with_fields(doc, headings):
    """Generate TOC with field codes for automatic page numbers"""
    # Add TOC title
    toc_heading = doc.add_paragraph()
    toc_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = toc_heading.add_run("Table of Contents")
    set_run_font(run, font_size=16, bold=True)
    
    # Add blank line
    doc.add_paragraph()
    
    # Add TOC entries with page number fields
    for heading in headings:
        toc_para = doc.add_paragraph()
        toc_para.paragraph_format.left_indent = Cm((heading['level'] - 1) * 0.5)
        
        # Add heading text
        run = toc_para.add_run(heading['text'])
        set_run_font(run, font_size=11)
        
        # Add tab leader (dots)
        run = toc_para.add_run(" .......... ")
        set_run_font(run, font_size=11)
        
        # Add PAGE field for page number
        run = toc_para.add_run()
        fldChar = OxmlElement('w:fldChar')
        fldChar.set(qn('w:fldCharType'), 'begin')
        run._element.append(fldChar)
        
        instrText = OxmlElement('w:instrText')
        instrText.text = "PAGE"
        run._element.append(instrText)
        
        fldChar = OxmlElement('w:fldChar')
        fldChar.set(qn('w:fldCharType'), 'end')
        run._element.append(fldChar)
        
        set_run_font(run, font_size=11)

def insert_toc_at_beginning(doc):
    """Insert TOC at the beginning of the document"""
    # Store all original content
    original_paragraphs = []
    for para in doc.paragraphs:
        original_paragraphs.append(para)
    
    # Find all numbered headings
    headings = []
    for para in original_paragraphs:
        if is_numbered_heading(para):
            level = get_heading_level(para.text)
            if level:
                headings.append({
                    'text': para.text,
                    'level': level
                })
    
    # Clear the document
    for paragraph in doc.paragraphs:
        p = paragraph._element
        p.getparent().remove(p)
    
    # Generate TOC with field codes
    generate_toc_with_fields(doc, headings)
    
    # Add page break after TOC
    doc.add_page_break()
    
    # Restore original content
    for para in original_paragraphs:
        new_para = doc.add_paragraph()
        
        # Copy runs
        for run in para.runs:
            new_run = new_para.add_run(run.text)
            if run.font.name:
                new_run.font.name = run.font.name
            if run.font.size:
                new_run.font.size = run.font.size
            new_run.font.bold = run.font.bold
            new_run.font.italic = run.font.italic
        
        # Copy paragraph formatting
        if para.alignment:
            new_para.alignment = para.alignment
        if para.paragraph_format.line_spacing:
            new_para.paragraph_format.line_spacing = para.paragraph_format.line_spacing
        if para.paragraph_format.space_before:
            new_para.paragraph_format.space_before = para.paragraph_format.space_before
        if para.paragraph_format.space_after:
            new_para.paragraph_format.space_after = para.paragraph_format.space_after
        if para.paragraph_format.first_line_indent:
            new_para.paragraph_format.first_line_indent = para.paragraph_format.first_line_indent
    
    return len(headings)

def format_document(input_path, output_path, generate_toc=True):
    """Main function to format the document"""
    try:
        # Open the document
        doc = Document(input_path)
        print("✓ Document opened successfully")
        
        # Italicize 'et al' throughout the document
        print("→ Italicizing 'et al'...")
        for paragraph in doc.paragraphs:
            italicize_et_al(paragraph)
        print("✓ 'et al' italicized")
        
        # Add page numbers at bottom center
        print("→ Adding page numbers at bottom center...")
        add_page_numbers(doc)
        print("✓ Page numbers added")
        
        # Generate Table of Contents if requested
        if generate_toc:
            print("→ Generating Table of Contents...")
            heading_count = insert_toc_at_beginning(doc)
            print(f"✓ Generated Table of Contents with {heading_count} headings")
        
        # Save the formatted document
        doc.save(output_path)
        
        print(f"\n{'='*60}")
        print(f"✓ DOCUMENT SUCCESSFULLY FORMATTED")
        print(f"✓ Saved to: {output_path}")
        print(f"{'='*60}")
        print("\n📋 INSTRUCTIONS FOR WPS OFFICE:")
        print("   1. Open the document in WPS Office")
        print("   2. Page numbers will appear at the bottom center")
        print("   3. To update Table of Contents page numbers:")
        print("      - Click anywhere in the Table of Contents")
        print("      - Press F9 on your keyboard")
        print("      - Select 'Update entire table'")
        print("      - Click OK")
        print("   4. The correct page numbers will appear automatically")
        print(f"{'='*60}\n")
        
    except Exception as e:
        print(f"✗ Error formatting document: {str(e)}")
        import traceback
        traceback.print_exc()
        sys.exit(1)

def main():
    """Main function to handle command line arguments"""
    if len(sys.argv) < 3 or len(sys.argv) > 4:
        print("="*60)
        print("WORD DOCUMENT FORMATTING SCRIPT")
        print("="*60)
        print("\nUsage: python format_word.py input.docx output.docx [--no-toc]")
        print("\nOptions:")
        print("  --no-toc    Skip table of contents generation")
        print("\nWhat this script does:")
        print("  1. Italicizes all 'et al' instances")
        print("  2. Adds page numbers at bottom center of every page")
        print("  3. Generates a Table of Contents with automatic page numbers")
        print("\nExample:")
        print("  python format_word.py my_document.docx formatted_document.docx")
        print("  python format_word.py my_document.docx formatted_document.docx --no-toc")
        print("\n" + "="*60)
        sys.exit(1)
    
    input_file = sys.argv[1]
    output_file = sys.argv[2]
    
    # Check for optional --no-toc flag
    generate_toc = True
    if len(sys.argv) > 3 and sys.argv[3] == '--no-toc':
        generate_toc = False
        print("⚠ Table of Contents generation disabled\n")
    
    # Check if input file exists
    if not os.path.exists(input_file):
        print(f"✗ Error: Input file '{input_file}' not found.")
        sys.exit(1)
    
    # Format the document
    format_document(input_file, output_file, generate_toc)

if __name__ == "__main__":
    main()
