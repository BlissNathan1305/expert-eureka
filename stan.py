import re
import sys
import os
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn


def format_run(run, bold=False, italic=False):
    """
    Apply consistent font styling: Times New Roman, 12pt
    """
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    run.bold = bold
    run.italic = italic

    # Extra enforcement (Word can be stubborn)
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')


def apply_inline_formatting(paragraph):
    """
    - Bold full subheading lines starting with numbering (e.g., 1.1 Title)
    - Italicize 'et al.'
    - Apply consistent font everywhere
    """
    text = paragraph.text.strip()

    if not text:
        return

    paragraph.clear()

    # Pattern for subheadings like 1.1, 2.4.2, etc.
    heading_pattern = r"^\d+(?:\.\d+)+\s+.+"

    # If it's a heading → bold everything
    if re.match(heading_pattern, text):
        run = paragraph.add_run(text)
        format_run(run, bold=True)
        return

    # Otherwise process "et al."
    parts = re.split(r"(et al\.)", text)

    for part in parts:
        if not part:
            continue

        if part == "et al.":
            run = paragraph.add_run(part)
            format_run(run, italic=True)
        else:
            run = paragraph.add_run(part)
            format_run(run)


def clean_docx(input_file, output_file):
    doc = Document(input_file)

    # Citation pattern
    bracket_pattern = re.compile(
        r"\[\s*\d+(?:\s*[-–]\s*\d+)?(?:\s*,\s*\d+(?:\s*[-–]\s*\d+)?)*\s*\]"
    )

    for paragraph in doc.paragraphs:
        original_text = paragraph.text

        # Remove citations
        cleaned_text = bracket_pattern.sub("", original_text)

        # Fix spacing issues
        cleaned_text = re.sub(r"\s{2,}", " ", cleaned_text)
        cleaned_text = re.sub(r"\s+([.,;:])", r"\1", cleaned_text)

        if cleaned_text.strip():
            paragraph.text = cleaned_text.strip()
            apply_inline_formatting(paragraph)

    doc.save(output_file)


def main():
    if len(sys.argv) != 3:
        print("Usage: python clean_docx.py <input.docx> <output.docx>")
        sys.exit(1)

    input_file = sys.argv[1]
    output_file = sys.argv[2]

    if not os.path.exists(input_file):
        print(f"Error: File '{input_file}' not found.")
        sys.exit(1)

    if not input_file.endswith(".docx"):
        print("Error: Input file must be a .docx file.")
        sys.exit(1)

    if not output_file.endswith(".docx"):
        print("Error: Output file must be a .docx file.")
        sys.exit(1)

    clean_docx(input_file, output_file)
    print(f"Done. Cleaned file saved as '{output_file}'")


if __name__ == "__main__":
    main()
